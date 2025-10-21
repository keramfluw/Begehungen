import streamlit as st
import pandas as pd
from io import BytesIO
from datetime import datetime, date

st.set_page_config(page_title="Begehungs-App (PV/Technik) – V3.1", layout="wide")

# -------- Optional dependency: python-docx (for Blanko-Formular) --------
DOCX_OK = True
try:
    from docx import Document
    from docx.shared import Pt
except Exception as e:
    DOCX_OK = False
    DOCX_ERR = str(e)

def build_blank_form_docx(templates: dict) -> bytes:
    if not DOCX_OK:
        raise RuntimeError(f"'python-docx' fehlt: {DOCX_ERR}")
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    doc.add_heading('Blanko-Formular – Begehung Kundenanlage', level=1)
    doc.add_paragraph("Hinweis: Dieses Formular dient zur Erfassung vor Ort, ohne App.")

    doc.add_heading('A. Kundendaten & Objekt', level=2)
    for label in ["Kunde / Ansprechpartner*in","E-Mail","Telefon","Adresse","Stadt","PLZ","Bundesland","Liegenschaftsnummer","Techniker*in / Team","Datum"]:
        doc.add_paragraph(f"{label}: _________________________________")

    doc.add_heading('B. Varianten (bitte ankreuzen)', level=2)
    doc.add_paragraph("[ ] Bronze     [ ] Silber     [ ] Gold")

    doc.add_heading('C. Checkliste – Prüfpunkte', level=2)
    rows = []
    seen = set()
    for v in ["Bronze","Silber","Gold"]:
        for it in templates.get(v, []):
            key = (it["item_group"], it["item_text"])
            if key in seen:
                continue
            seen.add(key)
            rows.append((it["item_group"], it["item_text"]))

    groups = {}
    for g, t in rows:
        groups.setdefault(g, []).append(t)

    for g, items in groups.items():
        doc.add_heading(f"Gruppe: {g}", level=3)
        table = doc.add_table(rows=len(items)+1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Prüfpunkt"
        hdr[1].text = "Status (ok/offen/kritisch/n/a)"
        hdr[2].text = "Wert/Einheit"
        hdr[3].text = "Notizen"
        for i, t in enumerate(items, start=1):
            table.rows[i].cells[0].text = t
            table.rows[i].cells[1].text = "____"
            table.rows[i].cells[2].text = "____"
            table.rows[i].cells[3].text = ""

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

# ----------------------------
# Session state init
# ----------------------------
if "inspections" not in st.session_state:
    st.session_state.inspections = pd.DataFrame(columns=[
        "inspection_id","date","technician","customer_name","customer_email","customer_phone",
        "address","city","plz","bundesland","liegenschaftsnummer",
        "variant_combo","item_id","item_group","item_text","status","value","unit","notes"
    ])

# Checklist templates (wie V3)
if "templates" not in st.session_state:
    st.session_state.templates = {
        "Bronze": [
            {"item_group":"Allgemein","item_text":"Zugang Dachflächen / Sicherheit (Geländer, Anschlagpunkte)","unit":"","default":"offen"},
            {"item_group":"PV/Elektrik","item_text":"Zählerschrank Zustand & Reserven","unit":"","default":"offen"},
            {"item_group":"PV/Elektrik","item_text":"Netzverknüpfungspunkt (Hausanschluss, NH, SLS)","unit":"","default":"offen"},
            {"item_group":"Gebäude","item_text":"Dachaufbau / Statik plausibel (Sichtprüfung)","unit":"","default":"offen"},
            {"item_group":"Dokumente","item_text":"Fotos/Skizze Dach (Ausrichtung, Hindernisse)","unit":"","default":"offen"},
            {"item_group":"Allgemein","item_text":"Begehungsanmeldung & Schlüsselkoordination (über kk&t)","unit":"","default":"offen"},
            {"item_group":"Planung","item_text":"Größe/Leistung PV grob bestimmen (qm × 0,25 kWp/qm) – Schätzung","unit":"kWp","default":"offen"},
            {"item_group":"Ertrag","item_text":"Dachausrichtung dokumentieren (keine Ertragsprognose)","unit":"","default":"offen"},
            {"item_group":"Standorte","item_text":"Orte für Technik (Speicher, WR, Notstrom) – Vorschlag","unit":"","default":"offen"},
            {"item_group":"Schaltschränke","item_text":"Elektrische Schaltschränke begutachten (Allgemeinzähler Keller)","unit":"","default":"offen"},
            {"item_group":"Kommunikation","item_text":"Videocall-Nachbesprechung (Termin anbieten)","unit":"","default":"offen"},
            {"item_group":"Finanzen","item_text":"Hinweis auf mögliche WEG-Finanzierung geben","unit":"","default":"offen"},
            {"item_group":"Infrastruktur","item_text":"Stellplatzsituation/Wallbox/Lademanagement – Sichtprüfung","unit":"","default":"offen"},
        ],
        "Silber": [
            {"item_group":"PV/Elektrik","item_text":"Einspeisepunkt / Messkonzept (Vorprüfung)","unit":"","default":"offen"},
            {"item_group":"PV/Elektrik","item_text":"Leitungswege (Dach → Zählerschrank)","unit":"","default":"offen"},
            {"item_group":"Gebäude","item_text":"Dachhaut / Abdichtung (Material, Alter, Zustand)","unit":"","default":"offen"},
            {"item_group":"Gebäude","item_text":"Blitz-/Potenzialausgleich (Bestand)","unit":"","default":"offen"},
            {"item_group":"Dokumente","item_text":"Planauszüge, Fotos, Maße (Beleg)","unit":"","default":"offen"},
            {"item_group":"Allgemein","item_text":"Begehungsanmeldung & Schlüsselkoordination (über kk&t)","unit":"","default":"offen"},
            {"item_group":"Planung","item_text":"Größe/Leistung PV: Objektfoto + tatsächliche Dachflächenberechnung","unit":"kWp","default":"offen"},
            {"item_group":"Ertrag","item_text":"Ertragsprognose/Jahr (belegte Dachausrichtung)","unit":"kWh/a","default":"offen"},
            {"item_group":"Regulatorik","item_text":"Netzverträglichkeitsprüfung/Anschlussbegehren – zubuchbares Paket 'Anmeldung/Anfrage Netzbetreiber'","unit":"","default":"offen"},
            {"item_group":"Machbarkeit","item_text":"Technische Realisierbarkeit PV (Ausrichtung/Verschattung) – Bewertung","unit":"","default":"offen"},
            {"item_group":"Netz","item_text":"Einholung Netzverträglichkeitsprüfung/Netzanschlussbegehren – optional (zubuchbar)","unit":"","default":"offen"},
            {"item_group":"Statik","item_text":"Dachlasten prüfen & Dachzustand fotografisch festhalten (optional Drohne)","unit":"","default":"offen"},
            {"item_group":"Elektrik","item_text":"Spätere Leitungsführung festlegen – Vorplanung (sofern möglich)","unit":"","default":"offen"},
            {"item_group":"Brandschutz","item_text":"Brandabschottungen – Vorprüfung (sofern relevant)","unit":"","default":"offen"},
            {"item_group":"Standorte","item_text":"Orte für Technik (Speicher, WR, Notstrom) – Konkretisierung","unit":"","default":"offen"},
            {"item_group":"Schaltschränke","item_text":"Zählerschrank-Bewertung PLUS Kostenschätzung für Ertüchtigung","unit":"€","default":"offen"},
            {"item_group":"Zähler/Mieterstrom","item_text":"Ausstattung Wohnungszähler bewerten (ohne Zählerplatzsichtung) + optional Werbebrief Mieterstrom (mit Zustimmung)","unit":"","default":"offen"},
            {"item_group":"Verbräuche","item_text":"Bisherige Stromverbräuche erheben (nur Betriebsstrom)","unit":"kWh/a","default":"offen"},
            {"item_group":"Kosten","item_text":"Vorschlag Verwendung Strom (WP, Betriebsstrom, Mieterstrom/GGV, Wallboxen) – grobe Vision","unit":"","default":"offen"},
            {"item_group":"Wirtschaftlichkeit","item_text":"Amortisationsrechnung – grobe Systemschätzung","unit":"","default":"offen"},
            {"item_group":"Kommunikation","item_text":"Videocall-Nachbesprechung durchführen","unit":"","default":"offen"},
            {"item_group":"Finanzen","item_text":"Hinweis auf mögliche WEG-Finanzierung geben","unit":"","default":"offen"},
            {"item_group":"Angebote","item_text":"Bewertung Drittangebote – NUR Texte/Ausschreibungsunterlagen sichten","unit":"","default":"offen"},
            {"item_group":"Infrastruktur","item_text":"Stellplatzsituation/Wallbox/Lademanagement – Grobkonzept","unit":"","default":"offen"},
            {"item_group":"Speicher","item_text":"Integration bestehender Speicher – Grobkonzept","unit":"","default":"offen"},
        ],
        "Gold": [
            {"item_group":"PV/Elektrik","item_text":"String-Layout & Wechselrichter-Standort (Vorplanung)","unit":"","default":"offen"},
            {"item_group":"PV/Elektrik","item_text":"Lastgänge / Verbrauchsstruktur (sofern vorhanden)","unit":"","default":"offen"},
            {"item_group":"Systeme","item_text":"Speicher / Ladeinfrastruktur / WP: Machbarkeit & Schnittstellen","unit":"","default":"offen"},
            {"item_group":"Regulatorik","item_text":"Messkonzept (GGV/Mieterstrom) – Detailaufnahme","unit":"","default":"offen"},
            {"item_group":"Risiken","item_text":"Sonderpunkte: Statik-Red Flags, Brandschutz, Denkmalschutz","unit":"","default":"offen"},
            {"item_group":"Allgemein","item_text":"Begehungsanmeldung & Schlüsselkoordination (über kk&t)","unit":"","default":"offen"},
            {"item_group":"Planung","item_text":"Größe/Leistung PV: Drohnenaufnahmen + 3D-Aufnahme (inkl. Bronze+Silber)","unit":"kWp","default":"offen"},
            {"item_group":"Ertrag","item_text":"Ertragsprognose/Jahr (szenariobasiert)","unit":"kWh/a","default":"offen"},
            {"item_group":"Vertragsmodelle","item_text":"Dachpacht oder Contracting-Konzepte prüfen (Pflicht in Gold)","unit":"","default":"offen"},
            {"item_group":"Machbarkeit","item_text":"Technische Realisierbarkeit PV (Ausrichtung/Verschattung) – Bewertung","unit":"","default":"offen"},
            {"item_group":"Netz","item_text":"Einholung Netzverträglichkeitsprüfung & Netzanschlussbegehren – inkludiert","unit":"","default":"offen"},
            {"item_group":"Statik","item_text":"Dachlasten prüfen & Dachzustand bebildert (inkl. Drohne möglich)","unit":"","default":"offen"},
            {"item_group":"Elektrik","item_text":"Spätere Leitungsführung festlegen – Vorplanung verbindlich","unit":"","default":"offen"},
            {"item_group":"Brandschutz","item_text":"Brandabschottungen – Prüfung/Erfordernis dokumentieren","unit":"","default":"offen"},
            {"item_group":"Standorte","item_text":"Orte für Technik (Speicher, WR, Notstrom) – finale Vorschläge","unit":"","default":"offen"},
            {"item_group":"Schaltschränke","item_text":"Zählerschrank – Fotos, Bewertung, Kostenschätzung + Video Keller/Heizung","unit":"","default":"offen"},
            {"item_group":"Zähler/Mieterstrom","item_text":"Ausstattung Wohnungszähler bewerten + Mieterabfrage (vor Mieteranschreiben)","unit":"","default":"offen"},
            {"item_group":"Verbräuche","item_text":"Bisherige Stromverbräuche erheben (Mieter/Nutzer + Betriebsstrom)","unit":"kWh/a","default":"offen"},
            {"item_group":"Baustelle","item_text":"Kosten Gebäudeeinrüstung & Baustelleneinrichtung – Abschätzung","unit":"€","default":"offen"},
            {"item_group":"WP","item_text":"Wärmepumpen-Machbarkeitsbetrachtung","unit":"","default":"offen"},
            {"item_group":"Nutzung","item_text":"Vorschlag Verwendung Strom (WP, Betriebsstrom, Mieterstrom/GGV, Wallboxen) – Konzeptvorschlag","unit":"","default":"offen"},
            {"item_group":"Wirtschaftlichkeit","item_text":"Amortisationsrechnung (mit/ohne Finanzierung, WEG-tauglich)","unit":"","default":"offen"},
            {"item_group":"Kommunikation","item_text":"Nachbetrachtung in VC + ggf. lokale Medienarbeit initiieren","unit":"","default":"offen"},
            {"item_group":"Kommunikation","item_text":"Videocall-Nachbesprechung durchführen","unit":"","default":"offen"},
            {"item_group":"Finanzen","item_text":"Hinweis auf mögliche WEG-Finanzierung geben","unit":"","default":"offen"},
            {"item_group":"Angebote","item_text":"Bewertung Drittangebote – inkl. Auswertung","unit":"","default":"offen"},
            {"item_group":"Infrastruktur","item_text":"Stellplatz/Wallbox/Lademanagement – Konzeptvorschlag","unit":"","default":"offen"},
            {"item_group":"Speicher","item_text":"Integration bestehender Speicher – Konzeptvorschlag","unit":"","default":"offen"},
            {"item_group":"IT/Schnittstellen","item_text":"Technische Schnittstellen (Zähler, BMS, EVSE, SG-Ready, API) – Klärung","unit":"","default":"offen"},
            {"item_group":"Schall","item_text":"Schallprüfung (WP/WR/Trafo) – Erfordernis & Maßnahmen","unit":"","default":"offen"},
            {"item_group":"Genehmigungen","item_text":"Genehmigungsprüfung (Bau/Denkmalschutz/Sonderfälle)","unit":"","default":"offen"},
        ]
    }

# Musterkunde Default (nutzt date.today() statt datetime)
if "musterkunde" not in st.session_state:
    st.session_state.musterkunde = {
        "customer_name": "WEG Beispielstraße 12",
        "customer_email": "verwaltung@example.com",
        "customer_phone": "0761-123456",
        "address": "Beispielstraße 12",
        "city": "Freiburg",
        "plz": "79100",
        "bundesland": "BW",
        "liegenschaftsnummer": "LG-2025-001",
        "technician": "Team Süd – Max & Lea",
        "date": date.today(),  # safer default
        "variants": ["Bronze","Silber"],
        "prefill_values": {
            ("Planung","Größe/Leistung PV grob bestimmen (qm × 0,25 kWp/qm) – Schätzung"): ("30", "kWp"),
            ("Ertrag","Ertragsprognose/Jahr (belegte Dachausrichtung)"): ("28500", "kWh/a"),
            ("Schaltschränke","Zählerschrank-Bewertung PLUS Kostenschätzung für Ertüchtigung"): ("2500", "€"),
        }
    }

def new_id(prefix="INS"):
    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    return f"{prefix}-{ts}"

# Sidebar navigation
st.sidebar.title("Navigation")
page = st.sidebar.radio("Ansicht wählen", [
    "Neue Begehung",
    "Bestand hochladen (CSV)",
    "Checklisten bearbeiten",
    "Datenexport / Reporting",
    "Blanko-Formular",
    "Hilfe"
])

# ----------------------------
# Page: Neue Begehung
# ----------------------------
if page == "Neue Begehung":
    st.title("📋 Neue Begehung – Musterkunde geladen")
    mk = st.session_state.musterkunde

    with st.form("form_begehung", clear_on_submit=False):
        st.subheader("Kundendaten & Objekt")
        cols = st.columns(3)
        customer_name = cols[0].text_input("Kunde / Ansprechpartner*in", value=mk["customer_name"])
        customer_email = cols[1].text_input("E-Mail", value=mk["customer_email"])
        customer_phone = cols[2].text_input("Telefon", value=mk["customer_phone"])

        colsa = st.columns(5)
        address = colsa[0].text_input("Adresse", value=mk["address"])
        city = colsa[1].text_input("Stadt", value=mk["city"])
        plz = colsa[2].text_input("PLZ", value=mk["plz"])
        bundesland = colsa[3].text_input("Bundesland", value=mk["bundesland"])
        liegenschaftsnummer = colsa[4].text_input("Liegenschaftsnummer", value=mk["liegenschaftsnummer"])

        st.subheader("Begehung")
        cols2 = st.columns(3)
        date_val = cols2[0].date_input("Datum", value=mk["date"])
        technician = cols2[1].text_input("Techniker*in / Team", value=mk["technician"])
        variants = cols2[2].multiselect("Variante(n) (frei kombinierbar)", ["Bronze","Silber","Gold"], default=mk["variants"])
        st.caption("Musterkunde ist vorausgefüllt. Sie können alles überschreiben.")

        # Build checklist
        selected_templates = []
        for v in variants:
            selected_templates.extend(st.session_state.templates.get(v, []))

        seen = set()
        checklist_rows = []
        for item in selected_templates:
            key = (item["item_group"], item["item_text"])
            if key in seen:
                continue
            seen.add(key)
            val = ""
            unit = item.get("unit","")
            if key in mk["prefill_values"]:
                val, unit_override = mk["prefill_values"][key]
                if unit_override:
                    unit = unit_override
            checklist_rows.append({
                "item_group": item["item_group"],
                "item_text": item["item_text"],
                "status": item.get("default","offen"),
                "value": val,
                "unit": unit,
                "notes": ""
            })

        st.subheader("Checkliste (vorausgefüllt)")
        edited_df = st.data_editor(
            pd.DataFrame(checklist_rows),
            num_rows="dynamic",
            use_container_width=True,
            column_config={
                "item_group": st.column_config.TextColumn("Gruppe"),
                "item_text": st.column_config.TextColumn("Prüfpunkt"),
                "status": st.column_config.SelectboxColumn("Status", options=["ok","offen","kritisch","n/a"]),
                "value": st.column_config.TextColumn("Wert/Messung"),
                "unit": st.column_config.TextColumn("Einheit"),
                "notes": st.column_config.TextColumn("Notizen"),
            },
            hide_index=True
        )

        c1, c2 = st.columns(2)
        submitted = c1.form_submit_button("✅ Begehung speichern")
        reset_to_muster = c2.form_submit_button("↺ Musterkunde erneut laden")

        if reset_to_muster:
            st.rerun()

        if submitted:
            inspection_id = new_id()
            variant_combo = "+".join(variants) if variants else "keine"
            # normalize
            records = []
            for i, row in edited_df.iterrows():
                records.append({
                    "inspection_id": inspection_id,
                    "date": pd.to_datetime(date_val),
                    "technician": technician,
                    "customer_name": customer_name,
                    "customer_email": customer_email,
                    "customer_phone": customer_phone,
                    "address": address, "city": city, "plz": plz, "bundesland": bundesland,
                    "liegenschaftsnummer": liegenschaftsnummer,
                    "variant_combo": variant_combo,
                    "item_id": f"ITM-{i+1:03d}",
                    "item_group": row["item_group"],
                    "item_text": row["item_text"],
                    "status": row["status"],
                    "value": row["value"],
                    "unit": row["unit"],
                    "notes": row["notes"],
                })
            if records:
                df_add = pd.DataFrame.from_records(records)
                st.session_state.inspections = pd.concat([st.session_state.inspections, df_add], ignore_index=True)
                st.success(f"Begehung **{inspection_id}** gespeichert ({len(records)} Zeilen).")
                st.download_button("⬇️ CSV dieser Begehung", data=df_add.to_csv(index=False).encode("utf-8"),
                                   file_name=f"{inspection_id}.csv", mime="text/csv")

# ----------------------------
# CSV Upload
# ----------------------------
elif page == "Bestand hochladen (CSV)":
    st.title("📤 CSV hochladen & zusammenführen")
    st.write("Erwartete Spalten (mindestens): inspection_id,date,technician,customer_name,address,city,plz,bundesland,liegenschaftsnummer,variant_combo,item_id,item_group,item_text,status,value,unit,notes")
    file = st.file_uploader("CSV-Datei wählen", type=["csv"])
    if file is not None:
        try:
            df_up = pd.read_csv(file)
            if "date" in df_up.columns:
                df_up["date"] = pd.to_datetime(df_up["date"], errors="coerce")
            st.dataframe(df_up.head(), use_container_width=True)
            if st.button("🔗 In Bestand übernehmen"):
                st.session_state.inspections = pd.concat([st.session_state.inspections, df_up], ignore_index=True).drop_duplicates()
                st.success(f"{len(df_up)} Zeilen übernommen.")
        except Exception as e:
            st.error(f"Fehler beim Einlesen: {e}")

# ----------------------------
# Templates bearbeiten
# ----------------------------
elif page == "Checklisten bearbeiten":
    st.title("🧩 Checklisten-Vorlagen je Variante")
    variants_all = list(st.session_state.templates.keys())
    selected = st.selectbox("Variante wählen", variants_all, index=0)
    df_tmpl = pd.DataFrame(st.session_state.templates[selected])
    edited = st.data_editor(
        df_tmpl,
        num_rows="dynamic",
        use_container_width=True,
        column_config={
            "item_group": st.column_config.TextColumn("Gruppe"),
            "item_text": st.column_config.TextColumn("Prüfpunkt"),
            "unit": st.column_config.TextColumn("Einheit"),
            "default": st.column_config.SelectboxColumn("Default-Status", options=["offen","ok","kritisch","n/a"]),
        },
        hide_index=True
    )
    if st.button("💾 Vorlage speichern"):
        st.session_state.templates[selected] = edited.to_dict(orient="records")
        st.success("Vorlage aktualisiert.")
    st.download_button("⬇️ Vorlage als CSV", data=edited.to_csv(index=False).encode("utf-8"),
                       file_name=f"vorlage_{selected.lower()}.csv", mime="text/csv")

# ----------------------------
# Export
# ----------------------------
elif page == "Datenexport / Reporting":
    st.title("📦 Export & Reporting")
    df = st.session_state.inspections.copy()
    if df.empty:
        st.info("Noch keine Daten vorhanden.")
    else:
        colf = st.columns(4)
        tech_filter = colf[0].text_input("Filter Techniker*in enthält")
        city_filter = colf[1].text_input("Filter Stadt enthält")
        status_filter = colf[2].selectbox("Filter Status", ["(alle)","ok","offen","kritisch","n/a"], index=0)
        variant_filter = colf[3].text_input("Filter Varianten enthalten (z. B. Bronze+Gold)")

        mask = pd.Series([True]*len(df))
        if tech_filter:
            mask &= df["technician"].fillna("").str.contains(tech_filter, case=False, regex=False)
        if city_filter:
            mask &= df["city"].fillna("").str.contains(city_filter, case=False, regex=False)
        if status_filter != "(alle)":
            mask &= df["status"].fillna("") == status_filter
        if variant_filter:
            mask &= df["variant_combo"].fillna("").str.contains(variant_filter, case=False, regex=False)

        view = df.loc[mask].sort_values(["date","inspection_id","item_id"])
        st.write(f"**{len(view)}** Zeilen im Filter")
        st.dataframe(view, use_container_width=True, height=400)

        def to_xlsx_bytes(df_export: pd.DataFrame) -> bytes:
            output = BytesIO()
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                df_export.to_excel(writer, index=False, sheet_name="Begehungen")
            return output.getvalue()

        st.download_button("⬇️ CSV", data=view.to_csv(index=False).encode("utf-8"),
                           file_name="begehungen_gefiltert.csv", mime="text/csv")
        st.download_button("⬇️ XLSX", data=to_xlsx_bytes(view),
                           file_name="begehungen_gefiltert.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ----------------------------
# Blanko-Formular
# ----------------------------
elif page == "Blanko-Formular":
    st.title("🖨️ Blanko-Formular zum Ausdrucken (DOCX)")
    if not DOCX_OK:
        st.error(f"Blanko-Formular benötigt 'python-docx'. Grund: {DOCX_ERR}")
        st.code("pip install python-docx")
    else:
        if st.button("📄 Blanko-Formular erzeugen"):
            try:
                doc_bytes = build_blank_form_docx(st.session_state.templates)
                st.download_button("⬇️ Blanko-Formular (DOCX)", data=doc_bytes, file_name="Blanko_Formular_Begehung.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"Fehler beim Erzeugen des Formulars: {e}")

# ----------------------------
# Hilfe
# ----------------------------
elif page == "Hilfe":
    st.title("ℹ️ Hilfe & Troubleshooting")
    st.markdown("""
**Neu in V3.1**
- Sicherere Datumsvorbelegung (`date.today()` statt `datetime`).
- Blanko-Formular mit klarer Fehlermeldung, wenn `python-docx` fehlt.
- `st.rerun()` statt `experimental_rerun`.

**Troubleshooting Quick-Checks**
1. Abhängigkeiten installiert? `pip install -r requirements.txt`
2. Python-Version 3.9–3.12 empfohlen.
3. Startbefehl: `streamlit run app.py`
4. Bei *ModuleNotFoundError* zu `python-docx`: `pip install python-docx`
5. Browser-Cache leeren oder Inkognito testen.
    """)